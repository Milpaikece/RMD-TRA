"""
RMD-TRA — pillar5/word_exporter.py
Generator laporan audit .docx menggunakan python-docx

Menghasilkan dua file Word:
1. Laporan Audit Lengkap + Solusi Revisi Siap Tempel
2. Draf IMRAD Scopus Q1 Siap Kirim
"""

from __future__ import annotations
import re
import io
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION, WD_ORIENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- LaTeX -> OMML (persamaan Word tergambar grafis) ---
# Opsional: kalau latex2mathml + XSLT Office tersedia, rumus $$...$$ dirender
# sebagai persamaan Word asli. Kalau tidak, otomatis fallback ke kode LaTeX.
try:
    import glob as _glob
    import latex2mathml.converter as _l2m
    from lxml import etree as _etree
    _OMML_LIBS = True
except ImportError:
    _OMML_LIBS = False

_OMML_XSLT = None  # None=belum dicoba, False=tak tersedia, else=objek transform


def _get_omml_xslt():
    """Muat MML2OMML.XSL bawaan Office (lazy, cari di lokasi umum)."""
    global _OMML_XSLT
    if _OMML_XSLT is not None:
        return _OMML_XSLT
    if not _OMML_LIBS:
        _OMML_XSLT = False
        return False
    patterns = [
        r"C:\Program Files\Microsoft Office\root\Office*\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office*\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office*\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\Office*\MML2OMML.XSL",
    ]
    for pat in patterns:
        for path in _glob.glob(pat):
            try:
                _OMML_XSLT = _etree.XSLT(_etree.parse(path))
                return _OMML_XSLT
            except Exception:
                continue
    _OMML_XSLT = False
    return False


def _latex_to_omml_element(latex: str):
    """LaTeX -> elemen OMML (<m:oMath>). None kalau konversi tak tersedia/gagal."""
    xslt = _get_omml_xslt()
    if not xslt:
        return None
    try:
        mathml = _l2m.convert(latex)
        omml = xslt(_etree.fromstring(mathml))
        return omml.getroot()
    except Exception:
        return None


def _check_docx():
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx belum terinstall. Jalankan: uv add python-docx"
        )


# ===========================================================================
# MESIN RENDER RICH-TEXT — konversi output mentah LLM (markdown + [TABEL])
# jadi format Word yang rapi, dipakai oleh semua bagian audit.
# ===========================================================================
_NAVY = RGBColor(0x0D, 0x1B, 0x2A)
_TEAL = RGBColor(0x1A, 0x7A, 0x6E)
_GREY = RGBColor(0x44, 0x44, 0x41)
_BORDER = RGBColor(0xD3, 0xD1, 0xC7)
# urutan penting: $...$ (rumus inline) dicek sebelum penanda lain
_INLINE_MD = re.compile(r'(\$[^$\n]+?\$|\*\*.+?\*\*|__.+?__|\*[^*\n]+?\*|`[^`\n]+?`)')


def _run(paragraph, text: str, size: float = 11, color=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def _add_inline_md(paragraph, text: str, size: float = 11, color=None) -> None:
    """Render **tebal**, *miring*, `kode`, dan $rumus$ inline; sisanya biasa."""
    pos = 0
    for m in _INLINE_MD.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], size, color)
        tok = m.group(0)
        if tok.startswith("$"):
            # rumus inline: tampilkan tanpa tanda $ dalam font matematika
            r = _run(paragraph, tok[1:-1], size, color)
            r.font.name = "Cambria Math"
        elif tok.startswith("**") or tok.startswith("__"):
            _run(paragraph, tok[2:-2], size, color).bold = True
        elif tok.startswith("`"):
            r = _run(paragraph, tok[1:-1], size, color)
            r.font.name = "Consolas"
        else:  # *miring*
            _run(paragraph, tok[1:-1], size, color).italic = True
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], size, color)


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _start_landscape(doc: "Document") -> None:
    """Buka seksi baru berorientasi landscape (untuk tabel lebar)."""
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    sec.left_margin = sec.right_margin = Cm(1.85)


def _end_landscape(doc: "Document") -> None:
    """Tutup seksi landscape, kembali ke A4 tegak seperti setelan laporan."""
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)


def _pack_row_cells(row: list[str]) -> list[str]:
    """Rapatkan sel berisi ke kiri, buang sel kosong & duplikat berurutan."""
    out: list[str] = []
    for c in row:
        c = c.strip()
        if c and (not out or out[-1] != c):
            out.append(c)
    return out


def _compact_sparse_table(grid: list[list[str]]) -> list[list[str]] | None:
    """
    Sebagian agen mereproduksi tabel spesifikasi (label: nilai) dari skripsi
    asli sebagai tabel markdown lebar dengan header terduplikasi dan tiap
    baris data hanya berisi 1-3 sel dari belasan kolom (sisanya kosong) —
    contoh nyata: baris "| | | Kapasitas Battery | 357 | Kwh | | | ... |"
    di tabel 15 kolom. Membuang kolom yang kosong DI SEMUA BARIS tidak
    menolong karena tiap kolom terisi di baris yang berbeda-beda.

    Fix: rapatkan sel berisi tiap baris ke kiri dan pangkas jumlah kolom ke
    yang benar-benar dipakai. Hanya diterapkan bila tabel memang jarang
    terisi (kepadatan rendah) supaya tabel yang sudah rapi tidak disentuh.
    Return None kalau tabel bukan kandidat (biarkan alur normal jalan).
    """
    if len(grid) < 2:
        return None
    ncol = len(grid[0])
    if ncol <= 6:
        return None
    data_rows = grid[1:]
    if not data_rows:
        return None
    density = sum(1 for row in data_rows for c in row if c) / (len(data_rows) * ncol)
    if density >= 0.35:
        return None

    packed_data = [_pack_row_cells(r) for r in data_rows]
    packed_data = [r for r in packed_data if r]
    if not packed_data:
        return None
    new_ncol = max(len(r) for r in packed_data)
    if new_ncol >= ncol:
        return None
    packed_data = [r + [""] * (new_ncol - len(r)) for r in packed_data]

    header = _pack_row_cells(grid[0])
    if len(header) <= 1:
        title = header[0] if header else ""
        header_row = [title] * new_ncol
    else:
        header_row = (header + [""] * new_ncol)[:new_ncol]
    return [header_row] + packed_data


def _render_table_from_rows(doc: "Document", rows: list[list[str]]) -> None:
    """
    Render list baris (list sel) jadi tabel Word ber-grid dengan header navy.

    Tabel dari dokumen Word asli sering membawa artefak merged-cell: sel
    gabungan terbaca berulang ("Depo Bus | Depo Bus", "D = 19,4 km" x10),
    baris header terduplikasi, dan kolom/sel kosong bertebaran — tampil
    berantakan kalau di-render apa adanya. Dibersihkan di sini:
    1) baris duplikat berurutan dibuang (artefak merge vertikal),
    2) tabel spesifikasi jarang-terisi dirapatkan ke kolom minimum,
    3) kolom yang kosong di SEMUA baris dibuang,
    4) sel identik berurutan digabung jadi satu sel lebar (merge horizontal),
    5) tabel lebar (>8 kolom) pakai font lebih kecil + autofit.
    """
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    grid = [[(r[j].strip() if j < len(r) else "") for j in range(ncol)] for r in rows]

    # 1) buang baris yang persis sama dengan baris sebelumnya
    dedup: list[list[str]] = []
    for r in grid:
        if dedup and r == dedup[-1]:
            continue
        dedup.append(r)
    grid = dedup

    # 2) rapatkan tabel spesifikasi label:nilai yang jarang terisi
    compacted = _compact_sparse_table(grid)
    if compacted is not None:
        grid = compacted
    ncol = len(grid[0]) if grid else ncol

    # 3) buang kolom yang kosong di semua baris
    keep = [j for j in range(ncol) if any(row[j] for row in grid)]
    if keep and len(keep) < ncol:
        grid = [[row[j] for j in keep] for row in grid]
    ncol = len(keep) if keep else ncol
    if ncol == 0:
        return

    # Tabel lebar tak muat di A4 tegak (teks patah vertikal). Diputar ke
    # halaman LANDSCAPE (lebar ~26cm) supaya kolom cukup lebar & terbaca.
    wide = ncol > 8
    if wide:
        _start_landscape(doc)
    size = 7 if ncol > 12 else (7.5 if wide else 9)

    tbl = doc.add_table(rows=len(grid), cols=ncol)
    tbl.style = "Table Grid"

    # lebar kolom PROPORSIONAL terhadap isi (bukan dibagi rata) — kolom
    # keterangan panjang dapat ruang lebar, kolom angka pendek jadi sempit.
    # Tanpa ini, tabel 13+ kolom dibagi rata dan teks patah vertikal
    # ("Batt ery", "CAPE X") sehingga tak terbaca.
    tbl.autofit = False
    usable_cm = 26.0 if wide else 15.5  # landscape 29,7 - 2x1,85 margin; portrait A4
    weights = [
        max(3.0, min(40.0, max(len(row[j]) for row in grid) or 1))
        for j in range(ncol)
    ]
    total_w = sum(weights)
    col_widths = [Cm(usable_cm * w / total_w) for w in weights]

    for i, r in enumerate(grid):
        j = 0
        while j < ncol:
            # rentang sel identik berurutan -> gabung jadi satu sel
            k = j
            while k + 1 < ncol and r[k + 1] == r[j] and r[j] != "":
                k += 1
            cell = tbl.rows[i].cells[j]
            if k > j:
                cell = cell.merge(tbl.rows[i].cells[k])
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline_md(para, r[j], size,
                           RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else _GREY)
            if i == 0:
                for run in para.runs:
                    run.bold = True
                _shade_cell(cell, "0D1B2A")
            j = k + 1

    # terapkan lebar per sel grid (merged cell otomatis menjumlahkan span-nya)
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            if j < ncol:
                cell.width = col_widths[j]
    doc.add_paragraph()

    if wide:
        _end_landscape(doc)


def _is_md_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r':?-{2,}:?', c) for c in cells if c != "")


def _extract_display_latex(s: str):
    """Kalau baris ini rumus display LaTeX ($$...$$ atau \\[...\\]), kembalikan isi LaTeX-nya; else None."""
    s = s.strip()
    m = re.fullmatch(r'\$\$(.+?)\$\$', s, re.DOTALL)
    if m:
        return m.group(1).strip()
    m = re.fullmatch(r'\\\[(.+?)\\\]', s, re.DOTALL)
    if m:
        return m.group(1).strip()
    return None


def _render_latex_formula(doc: "Document", latex: str) -> None:
    """
    Render rumus LaTeX. Diutamakan sebagai PERSAMAAN WORD ASLI (tergambar grafis
    via OMML). Kalau konversi tidak tersedia (mis. komputer tanpa Office/latex2mathml),
    fallback ke kode LaTeX terpusat font Cambria Math yang tetap rapi & bisa disalin.
    """
    omml = _latex_to_omml_element(latex)
    if omml is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p._p.append(omml)
        return

    # Fallback: tampilkan kode LaTeX
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    safe = "".join(ch for ch in latex if ch >= " " or ch in "\t")
    run = p.add_run(safe)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)


def _render_rich_text(doc: "Document", text: str, base_size: float = 11) -> None:
    """
    Konversi teks mentah (markdown + blok [TABEL]) jadi elemen Word rapi:
    heading, tabel asli, bullet, label terstruktur, dan inline bold/italic.
    """
    if not text or not text.strip():
        return

    # Pisahkan blok [TABEL]...[/TABEL] dulu, render bagian teks di antaranya.
    pos = 0
    for m in re.finditer(r'\[TABEL\](.*?)\[/TABEL\]', text, re.DOTALL):
        _render_text_lines(doc, text[pos:m.start()], base_size)
        rows = [
            [c.strip() for c in ln.strip().strip("|").split("|")]
            for ln in m.group(1).splitlines() if ln.strip()
        ]
        _render_table_from_rows(doc, rows)
        pos = m.end()
    _render_text_lines(doc, text[pos:], base_size)


def _render_text_lines(doc: "Document", text: str, base_size: float = 11) -> None:
    """Proses teks per baris: heading, separator, bullet, label, tabel markdown."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        # lewati baris kosong & pagar kode ```
        if not s or s.startswith("```"):
            i += 1
            continue

        # rumus LaTeX display: baris "$$ ... $$" atau "\[ ... \]" -> render
        # sebagai blok rumus terpusat & monospace (rapi + gampang di-copy)
        latex = _extract_display_latex(s)
        if latex is not None:
            _render_latex_formula(doc, latex)
            i += 1
            continue

        # tabel markdown: baris "|...|" berturut-turut (opsional baris pemisah |---|)
        if s.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_md_separator_row(lines[i]):
                    block.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _render_table_from_rows(doc, block)
            continue

        # heading markdown  #, ##, ###
        hm = re.match(r'^(#{1,6})\s+(.*)$', s)
        if hm:
            lvl = len(hm.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = _run(p, hm.group(2).replace("**", "").strip(),
                     12 if lvl <= 2 else 11, _TEAL if lvl <= 2 else _GREY)
            r.bold = True
            i += 1
            continue

        # garis pemisah (hanya karakter garis)
        if len(s) >= 3 and set(s) <= set("-─—=*_·• "):
            sep = doc.add_paragraph("─" * 50)
            for r in sep.runs:
                r.font.color.rgb = _BORDER
                r.font.size = Pt(8)
            i += 1
            continue

        # bullet  - / * / • / "1."
        bm = re.match(r'^\s*(?:[-*•]|\d+\.)\s+(.*)$', raw)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_md(p, bm.group(1), base_size)
            i += 1
            continue

        # label italic di awal baris "*Narasi:*", "*Teks Asli:*",
        # "*Solusi Revisi (siap tempel ke Word):*" -> BOLD + ITALIC, sisanya normal
        im = re.match(r'^\*([^*\n]*:)\*\s*(.*)$', s)
        if im:
            p = doc.add_paragraph()
            lr = _run(p, im.group(1), base_size, _NAVY)
            lr.bold = True
            lr.italic = True
            if im.group(2):
                _run(p, " ", base_size)
                _add_inline_md(p, im.group(2), base_size)
            i += 1
            continue

        # label terstruktur ALL-CAPS:  "LOKASI:", "GALAT:", "KOREKSI TEKNIS:"
        lm = re.match(r'^([A-Z][A-Z0-9 \-/()&]{1,45}):\s*(.*)$', s)
        if lm:
            p = doc.add_paragraph()
            _run(p, lm.group(1) + ": ", base_size, _NAVY).bold = True
            if lm.group(2):
                _add_inline_md(p, lm.group(2), base_size)
            i += 1
            continue

        # paragraf biasa
        p = doc.add_paragraph()
        _add_inline_md(p, s, base_size)
        i += 1


def _add_heading(doc: "Document", text: str, level: int = 1) -> None:
    """Tambah heading dengan style yang konsisten."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)  # Navy
            run.font.size = Pt(14)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x6E)  # Teal
            run.font.size = Pt(12)
        else:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x41)
            run.font.size = Pt(11)


def _add_perspective_label(doc: "Document", label: str) -> None:
    """
    Tandai bagian audit dengan sudut pandang penguji yang diwakilinya —
    4 sub-agen paralel (Consistency/Stats/Discussion/Ghost Citation) secara
    alami mewakili 4 fokus penguji berbeda, jadi cukup dilabeli eksplisit
    di sini daripada membuat agen "persona penguji" terpisah yang duplikat.
    """
    p = doc.add_paragraph()
    run = p.add_run(f"Perspektif Penguji: {label}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x65, 0x60)


def _add_info_box(doc: "Document", label: str, content: str,
                  color: str = "INFO") -> None:
    """Tambah kotak informasi berwarna untuk solusi revisi."""
    colors = {
        "KRITIS": RGBColor(0xC0, 0x39, 0x2B),   # Merah
        "SEDANG": RGBColor(0xBA, 0x75, 0x17),    # Amber
        "INFO":   RGBColor(0x18, 0x5F, 0xA5),    # Biru
        "SOLUSI": RGBColor(0x0F, 0x6E, 0x56),    # Hijau/Teal
    }
    para = doc.add_paragraph()
    label_run = para.add_run(f"[{label}] ")
    label_run.bold = True
    label_run.font.color.rgb = colors.get(color, colors["INFO"])
    label_run.font.size = Pt(11)
    content_run = para.add_run(content)
    content_run.font.size = Pt(11)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(6)


def _add_revision_block(doc: "Document", location: str,
                        original: str, revision: str) -> None:
    """Tambah blok solusi revisi dengan format tiga bagian."""
    doc.add_paragraph()
    _add_info_box(doc, "LOKASI", location, "INFO")

    if original:
        orig_para = doc.add_paragraph()
        orig_run = orig_para.add_run("Teks Asli: ")
        orig_run.bold = True
        orig_run.font.color.rgb = RGBColor(0xA3, 0x2D, 0x2D)
        orig_run.font.size = Pt(10)
        orig_text = orig_para.add_run(original)
        orig_text.font.size = Pt(10)
        orig_text.font.italic = True
        orig_para.paragraph_format.left_indent = Cm(1)

    rev_para = doc.add_paragraph()
    rev_run = rev_para.add_run("Teks Revisi (Siap Tempel): ")
    rev_run.bold = True
    rev_run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    rev_run.font.size = Pt(10)
    rev_text = rev_para.add_run(revision)
    rev_text.font.size = Pt(10)
    rev_para.paragraph_format.left_indent = Cm(1)

    # Garis pemisah
    sep = doc.add_paragraph("─" * 60)
    sep.paragraph_format.left_indent = Cm(0)
    for run in sep.runs:
        run.font.color.rgb = RGBColor(0xD3, 0xD1, 0xC7)
        run.font.size = Pt(8)


_SIMPUL_COLORS = {
    "TERHUBUNG": RGBColor(0x0F, 0x6E, 0x56),   # Hijau teal
    "LEMAH":     RGBColor(0xBA, 0x75, 0x17),    # Amber
    "TERPUTUS":  RGBColor(0xC0, 0x39, 0x2B),    # Merah
}

_SIMPUL_LABELS = [
    ("SIMPUL 1", "Bab 1 → Bab 2", "Rumusan Masalah → Tinjauan Literatur"),
    ("SIMPUL 2", "Bab 2 → Bab 3", "Teori → Metode yang Tepat"),
    ("SIMPUL 3", "Bab 3 → Bab 4", "Metode → Hasil & Pembahasan (merujuk Bab 2)"),
    ("SIMPUL 4", "Bab 4 → Bab 5", "Hasil → Sintesis Kesimpulan"),
    ("SIMPUL 5", "Bab 1 ↔ Bab 5", "Sirkularitas Rumusan Masalah ↔ Kesimpulan"),
]


def _add_simpul_table(doc: "Document", simpul_text: str) -> None:
    """
    Render tabel 5 SIMPUL benang merah sebagai tabel Word berwarna.
    Parsing format: SIMPUL N | ... | ... | [STATUS: ...]
    """
    # Parse baris status dari teks agen
    status_map: dict[int, str] = {}
    for line in simpul_text.splitlines():
        m = re.match(r'SIMPUL\s+(\d)\s*\|.*\[STATUS:\s*(TERHUBUNG|LEMAH|TERPUTUS)\]',
                     line, re.IGNORECASE)
        if m:
            status_map[int(m.group(1))] = m.group(2).upper()

    # Judul tabel
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Kerangka Evaluasi 5 SIMPUL Benang Merah — Alur Berpikir Ilmiah Skripsi"
    )
    intro_run.bold = True
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

    # Buat tabel: Simpul | Transisi | Yang Diperiksa | Status
    tbl = doc.add_table(rows=6, cols=4)
    tbl.style = "Table Grid"

    # Header
    headers = ["Simpul", "Transisi Antar Bab", "Yang Diperiksa", "Status"]
    hdr_row = tbl.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = hdr
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Warna header navy
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0D1B2A')
        tcPr.append(shd)

    # Baris data
    for i, (simpul_id, transisi, diperiksa) in enumerate(_SIMPUL_LABELS):
        row = tbl.rows[i + 1]
        status = status_map.get(i + 1, "—")
        data = [simpul_id, transisi, diperiksa, status]
        color = _SIMPUL_COLORS.get(status, RGBColor(0x44, 0x44, 0x44))

        for j, val in enumerate(data):
            cell = row.cells[j]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if j == 3:  # Kolom status
                        run.bold = True
                        run.font.color.rgb = color

    doc.add_paragraph()  # spasi setelah tabel


_MATRIX_STATUS_COLORS = {
    "TERJAWAB": RGBColor(0x0F, 0x6E, 0x56),        # hijau
    "PERLU DICEK": RGBColor(0xBA, 0x75, 0x17),      # amber
    "TIDAK TERJAWAB": RGBColor(0xC0, 0x39, 0x2B),   # merah
    # kompatibilitas label lama
    "LENGKAP": RGBColor(0x0F, 0x6E, 0x56),
    "TIDAK LENGKAP": RGBColor(0xC0, 0x39, 0x2B),
}


def _normalize_matrix_status(raw: str) -> str:
    """Ambil status dari sel terakhir matriks — terima nilai polos MAUPUN [STATUS: X]."""
    s = raw.strip()
    m = re.search(r'\[STATUS:\s*([^\]]+)\]', s, re.IGNORECASE)
    if m:
        s = m.group(1)
    s = s.strip().upper()
    for known in ("TIDAK TERJAWAB", "PERLU DICEK", "TERJAWAB", "TIDAK LENGKAP", "LENGKAP"):
        if known in s:
            return known
    return "—"


def _strip_machine_blocks(text: str) -> str:
    """
    Buang blok data mesin [TABEL SIMPUL] & [MATRIKS KETERLACAKAN] (beserta baris
    mentahnya "SIMPUL n | ... " / "RMn | ...") dari teks narasi/section, karena
    blok itu sudah dirender terpisah sebagai TABEL. Tanpa ini, baris pipa mentah
    ikut muncul sebagai paragraf jelek di laporan (bocor dua kali).
    """
    text = re.sub(r'\[TABEL SIMPUL\].*?\[/TABEL SIMPUL\]', '', text,
                  flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'\[MATRIKS KETERLACAKAN\].*?\[/MATRIKS KETERLACAKAN\]', '', text,
                  flags=re.DOTALL | re.IGNORECASE)
    # buang marker yatim (blok tanpa pasangan) dan baris mesin yang tercecer
    text = re.sub(r'\[/?(?:TABEL SIMPUL|MATRIKS KETERLACAKAN)\]', '', text,
                  flags=re.IGNORECASE)
    text = re.sub(r'(?m)^\s*(?:SIMPUL\s*\d|RM\d)\s*\|.*$', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    return text


def _add_traceability_matrix(doc: "Document", matrix_text: str) -> None:
    """
    Render tabel JAWABAN Rumusan Masalah dengan skema analitis:
    RM | Rumusan Masalah | Hasil Penelitian | Pembahasan | Status.
      - Hasil Penelitian : temuan faktual + angka kunci (narasi mengalir).
      - Pembahasan       : ANALISIS — makna → kontras kekuatan/kelemahan →
                           implikasi/arah (bukan sekadar ulang angka).
      - Status           : TERJAWAB / PERLU DICEK / TIDAK TERJAWAB.
    """
    rows_data = []
    for line in matrix_text.splitlines():
        line = line.strip()
        if not line or not line.upper().startswith("RM"):
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 5:
            continue
        status = _normalize_matrix_status(parts[-1])
        rm, rumusan = parts[0], parts[1]
        if len(parts) >= 6:
            # Format lengkap: RM|Rumusan|Hasil Penelitian|Pembahasan|Kesimpulan|Status
            hasil, pembahasan, kesimpulan = parts[2], parts[3], parts[4]
        else:
            # Format 5 kolom (tanpa Kesimpulan) — toleransi, Kesimpulan kosong.
            hasil, pembahasan, kesimpulan = parts[2], parts[3], ""
        # buang label redundan di awal sel (header kolom sudah menyebutnya)
        hasil = re.sub(r'^\s*(?:Hasil Penelitian|TEMUAN)\s*:\s*', '', hasil, flags=re.I)
        pembahasan = re.sub(r'^\s*(?:Pembahasan|MAKNA)\s*:\s*', '', pembahasan, flags=re.I)
        kesimpulan = re.sub(r'^\s*Kesimpulan\s*:\s*', '', kesimpulan, flags=re.I)
        rows_data.append((rm, rumusan, hasil, pembahasan, kesimpulan, status))

    if not rows_data:
        return  # tidak ada data valid, jangan tampilkan tabel kosong

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Jawaban Rumusan Masalah — Hasil Penelitian, Pembahasan & Kesimpulan (analitis)"
    )
    intro_run.bold = True
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

    # Tiga kolom naratif (Hasil, Pembahasan, Kesimpulan) butuh ruang — pakai
    # orientasi landscape supaya tidak berdesakan / teks tidak patah-patah.
    _start_landscape(doc)

    tbl = doc.add_table(rows=len(rows_data) + 1, cols=6)
    tbl.style = "Table Grid"
    tbl.autofit = False
    # lebar (A4 landscape, usable ~26cm)
    col_w = [Cm(1.1), Cm(3.6), Cm(7.0), Cm(7.0), Cm(5.5), Cm(1.8)]

    headers = ["RM", "Rumusan Masalah", "Hasil Penelitian", "Pembahasan",
               "Kesimpulan", "Status"]
    hdr_row = tbl.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = hdr
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "0D1B2A")

    STATUS_COL = 5
    for i, row_vals in enumerate(rows_data):
        row = tbl.rows[i + 1]
        color = _MATRIX_STATUS_COLORS.get(row_vals[STATUS_COL], RGBColor(0x44, 0x44, 0x44))
        for j, val in enumerate(row_vals):
            cell = row.cells[j]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if j == STATUS_COL:
                        run.bold = True
                        run.font.color.rgb = color

    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            if j < len(col_w):
                cell.width = col_w[j]

    _end_landscape(doc)
    doc.add_paragraph()


def _add_simpul_revision_blocks(doc: "Document", consistency_text: str) -> None:
    """
    Parse dan render blok [REVISI SIMPUL N]...[/REVISI SIMPUL]
    dari output Consistency Engine.
    """
    blocks = re.findall(
        r'\[REVISI SIMPUL (\d)\](.*?)\[/REVISI SIMPUL\]',
        consistency_text, re.DOTALL | re.IGNORECASE
    )
    if not blocks:
        # Fallback: tampilkan teks mentah jika format baru belum ada
        _render_rich_text(doc, consistency_text)
        return

    for simpul_num, block_text in blocks:
        # Sub-heading per simpul
        h = doc.add_heading(f"Simpul {simpul_num} — Temuan & Revisi", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x6E)
            run.font.size = Pt(12)

        # Parse field dalam blok
        fields = {
            "TRANSISI":    r'TRANSISI:\s*(.*?)(?=LOKASI:|MASALAH:|TEKS ASLI:|TEKS REVISI:|$)',
            "LOKASI":      r'LOKASI:\s*(.*?)(?=TRANSISI:|MASALAH:|TEKS ASLI:|TEKS REVISI:|$)',
            "MASALAH":     r'MASALAH:\s*(.*?)(?=TRANSISI:|LOKASI:|TEKS ASLI:|TEKS REVISI:|$)',
            "TEKS ASLI":   r'TEKS ASLI:\s*(.*?)(?=TRANSISI:|LOKASI:|MASALAH:|TEKS REVISI:|$)',
            "TEKS REVISI": r'TEKS REVISI:\s*(.*?)(?=TRANSISI:|LOKASI:|MASALAH:|TEKS ASLI:|$)',
        }
        parsed: dict[str, str] = {}
        for fname, fpat in fields.items():
            m = re.search(fpat, block_text, re.DOTALL | re.IGNORECASE)
            parsed[fname] = m.group(1).strip() if m else ""

        if parsed.get("TRANSISI"):
            _add_info_box(doc, "Transisi", parsed["TRANSISI"], "INFO")
        if parsed.get("LOKASI"):
            _add_info_box(doc, "Lokasi", parsed["LOKASI"], "INFO")
        if parsed.get("MASALAH"):
            _add_info_box(doc, "Masalah Benang Merah", parsed["MASALAH"], "SEDANG")
        if parsed.get("TEKS ASLI") or parsed.get("TEKS REVISI"):
            _add_revision_block(
                doc,
                location=parsed.get("LOKASI", ""),
                original=parsed.get("TEKS ASLI", ""),
                revision=parsed.get("TEKS REVISI", "(Lihat narasi audit di atas)"),
            )


# Pola header start dibuat TOLERAN: agen kadang mengganti em-dash "—" dengan
# "-"/"–", atau menghilangkan kurung siku. Kuncinya frasa inti header saja.
# Marker penutup [X_DONE] jauh lebih stabil (token sederhana), jadi dipakai
# sebagai jangkar fallback kalau header start tetap tidak dikenali.
_D = r"[—–-]+"  # kelas tanda hubung
_SECTION_MARKERS = {
    "consistency":  (rf"\[?\s*SOLUSI REVISI\s*{_D}\s*CONSISTENCY ENGINE\s*\]?",
                     r"\[CONSISTENCY_DONE\]"),
    "stats":        (rf"\[?\s*SOLUSI REVISI\s*{_D}\s*STATISTICAL AUDITOR\s*\]?",
                     r"\[STATS_AUDITOR_DONE\]"),
    "discussion":   (rf"\[?\s*SOLUSI REVISI\s*{_D}\s*DISCUSSION CRITIQUE\s*\]?",
                     r"\[DISCUSSION_DONE\]"),
    "ghost":        (rf"\[?\s*SOLUSI REVISI\s*{_D}\s*GHOST CITATION DETECTOR\s*\]?",
                     r"\[GHOST_CITATION_DONE\]"),
    "journal":      (rf"\[?\s*DRAF IMRAD\s*{_D}\s*SIAP TEMPEL[^\]\n]*\]?",
                     r"\[JOURNAL_DONE\]"),
    "draf_revisi_id": (rf"\[?\s*DRAF REVISI PENUH\s*{_D}\s*BAHASA INDONESIA\s*\]?",
                     r"\[JOURNAL_DONE\]"),
    "paraphrase":   (rf"\[?\s*PARAFRASE SIAP TEMPEL\s*{_D}\s*PARAPHRASING ENGINE\s*\]?",
                     r"\[PARAPHRASE_DONE\]"),
    "template":     (rf"\[?\s*AUTO[- ]TEMPLATE CONVERTER[^\]\n]*\]?",
                     r"\[TEMPLATE_DONE\]"),
    "sidang":       (rf"\[?\s*SIMULASI SIDANG\s*{_D}\s*PERTANYAAN\s*&\s*JAWABAN\s*\]?",
                     r"\[SIDANG_PREP_DONE\]"),
}
# draf_revisi_id berbagi marker penutup dengan journal — fallback berbasis
# penutup akan salah ambil isi journal, jadi khusus dia tanpa fallback.
_NO_FALLBACK = {"draf_revisi_id"}


def _parse_sections(audit_text: str) -> dict[str, str]:
    """
    Pisahkan output gabungan dari semua sub-agen berdasarkan marker.
    Dua lapis: (1) header start toleran; (2) kalau header tetap tak dikenali
    tapi marker penutup [X_DONE] ada, ambil teks sejak marker penutup section
    sebelumnya — supaya konten agen tidak pernah hilang dari laporan hanya
    karena agen menulis header-nya sedikit berbeda.
    """
    # kumpulkan posisi SEMUA marker (header start maupun penutup) sebagai
    # jangkar fallback dua arah
    end_positions = []
    all_marker_positions = []
    for _, (start_pat, end_pat) in _SECTION_MARKERS.items():
        for m in re.finditer(end_pat, audit_text, re.IGNORECASE):
            end_positions.append(m.end())
            all_marker_positions.append(m.start())
        for m in re.finditer(start_pat, audit_text, re.IGNORECASE):
            all_marker_positions.append(m.start())
    end_positions.sort()
    all_marker_positions.sort()

    sections = {}
    for key, (start_pat, end_pat) in _SECTION_MARKERS.items():
        match = re.search(
            f"{start_pat}(.*?){end_pat}",
            audit_text, re.DOTALL | re.IGNORECASE
        )
        if match:
            sections[key] = match.group(1).strip()
            continue
        if key in _NO_FALLBACK:
            sections[key] = ""
            continue

        start_m = re.search(start_pat, audit_text, re.IGNORECASE)
        end_m = re.search(end_pat, audit_text, re.IGNORECASE)

        if start_m and not end_m:
            # Fallback A: header ada tapi agen lupa marker penutup — ambil
            # sampai marker section BERIKUTNYA (header/penutup apa pun), atau EOF.
            nxt = min((p for p in all_marker_positions if p >= start_m.end()),
                      default=len(audit_text))
            sections[key] = audit_text[start_m.end():nxt].strip()
        elif end_m:
            # Fallback B: penutup ada tapi header tak dikenali — ambil sejak
            # marker penutup section sebelumnya.
            prev_end = max((p for p in end_positions if p <= end_m.start()), default=0)
            sections[key] = audit_text[prev_end:end_m.start()].strip()
        else:
            sections[key] = ""
    return sections


def generate_audit_report(
    audit_text: str,
    thesis_title: str,
    student_name: str = "Mahasiswa",
    output_path: Optional[str] = None,
    finance_verification: Optional[str] = None,
    extraction_warnings: Optional[list] = None,
    survey_verification: Optional[str] = None,
) -> bytes:
    """
    Hasilkan file Word laporan audit lengkap dengan solusi revisi siap tempel.

    Args:
        audit_text  : Output gabungan dari semua 8 sub-agen.
        thesis_title: Judul skripsi yang dianalisis.
        student_name: Nama mahasiswa.
        output_path : Jika diisi, simpan ke path ini. Jika None, return bytes.
        finance_verification: Hasil verifikasi angka otomatis dari
            _build_finance_verification_block (Python, bukan LLM) — kalau
            diisi, ditampilkan sebagai bagian tersendiri di awal laporan.
        extraction_warnings: Daftar peringatan dari _diagnose_chapter_extraction
            (pembagian bab mencurigakan). Kalau ada, ditampilkan sebagai kotak
            peringatan MENCOLOK di paling atas laporan supaya auditor manusia
            tahu hasil audit mungkin tidak andal SEBELUM mempercayainya —
            mengubah kegagalan ekstraksi yang diam-diam jadi bersuara.

    Returns:
        bytes konten file .docx
    """
    _check_docx()
    doc = Document()

    # ── Pengaturan halaman A4 ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ── Halaman judul ──────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("LAPORAN AUDIT SKRIPSI")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("RMD-TRA — Research Management & Development")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x6E)

    doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Judul Skripsi", thesis_title),
        ("Nama Mahasiswa", student_name),
        ("Tanggal Audit", datetime.now().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()

    # ── Peringatan Sistem: pembagian bab mencurigakan ────────────────────
    # Ditaruh PALING ATAS (sebelum apa pun) supaya auditor manusia melihatnya
    # lebih dulu. Kalau ekstraksi bab meragukan, seluruh audit di bawahnya
    # bisa jadi menilai teks yang salah — jujur katakan itu, jangan diam.
    if extraction_warnings:
        warn_head = doc.add_paragraph()
        wh_run = warn_head.add_run("⚠ PERINGATAN SISTEM — Pembagian Bab Perlu Diperiksa")
        wh_run.bold = True
        wh_run.font.size = Pt(13)
        wh_run.font.color.rgb = RGBColor(0xB0, 0x2A, 0x1E)

        warn_intro = doc.add_paragraph()
        wi_run = warn_intro.add_run(
            "Sistem mendeteksi kejanggalan saat memisahkan bab skripsi ini secara "
            "otomatis. Karena setiap sub-agen menilai teks per bab, pembagian yang "
            "salah dapat membuat sebagian temuan di bawah menjadi tidak akurat. "
            "Periksa poin berikut secara manual sebelum mempercayai isi laporan:"
        )
        wi_run.font.size = Pt(10)
        wi_run.font.italic = True
        wi_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        for w in extraction_warnings:
            wp = doc.add_paragraph(style="List Bullet")
            wr = wp.add_run(str(w))
            wr.font.size = Pt(10)
            wr.font.color.rgb = RGBColor(0xB0, 0x2A, 0x1E)
        doc.add_page_break()

    # ── Bagian 0: Verifikasi Angka Otomatis (deterministik, bukan LLM) ────
    if finance_verification and finance_verification.strip():
        _add_heading(doc, "0. Verifikasi Angka Otomatis (Sistem)", 1)
        notice_fv = doc.add_paragraph()
        notice_fv_run = notice_fv.add_run(
            "Bagian ini dihasilkan langsung oleh perhitungan Python — bukan LLM — "
            "sehingga dijamin akurat secara matematis. Dipakai sebagai bukti oleh "
            "Sub-Agen Statistical Auditor di Bagian 2."
        )
        notice_fv_run.font.italic = True
        notice_fv_run.font.size = Pt(10)
        notice_fv_run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        # buang garis pembatas "====" dan judul duplikat, sisanya render rapi
        fv_clean = "\n".join(
            ln for ln in finance_verification.splitlines()
            if ln.strip() and set(ln.strip()) != {"="}
            and not ln.strip().startswith("VERIFIKASI ANGKA OTOMATIS")
        )
        _render_rich_text(doc, fv_clean, base_size=10)
        doc.add_page_break()

    # ── Bagian 0b: Verifikasi Survei Otomatis (Slovin, %, rasio) ──────────
    if survey_verification and survey_verification.strip():
        _add_heading(doc, "0b. Verifikasi Survei Otomatis (Sistem)", 1)
        notice_sv = doc.add_paragraph()
        notice_sv_run = notice_sv.add_run(
            "Bagian ini juga dihitung langsung oleh Python — bukan LLM. Untuk "
            "skripsi berbasis survei/kualitas layanan, di sinilah ukuran sampel "
            "(rumus Slovin), jumlah persentase, dan klaim rasio diverifikasi ulang "
            "secara pasti."
        )
        notice_sv_run.font.italic = True
        notice_sv_run.font.size = Pt(10)
        notice_sv_run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        sv_clean = "\n".join(
            ln for ln in survey_verification.splitlines()
            if ln.strip() and set(ln.strip()) != {"="}
            and not ln.strip().startswith("VERIFIKASI SURVEI OTOMATIS")
        )
        _render_rich_text(doc, sv_clean, base_size=10)
        doc.add_page_break()

    # ── Parse sections dari output agen ──────────────────────────────────
    sections = _parse_sections(audit_text)

    # ── Bagian 1: Consistency Engine — Kerangka 5 SIMPUL ─────────────────
    _add_heading(doc, "1. Analisis Konsistensi Benang Merah (Bab 1–5)", 1)
    _add_perspective_label(doc, "Alur Berpikir Ilmiah & Keterlacakan Argumen")

    # Penjelasan singkat kerangka — selalu tampil sebagai fitur tetap
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(
        "Analisis ini mengevaluasi keutuhan benang merah (logical thread) yang "
        "menghubungkan seluruh bab skripsi secara sirkular. Kerangka 5 SIMPUL "
        "di bawah ini memetakan alur berpikir ilmiah: dari Rumusan Masalah (Bab 1) "
        "yang menuntun pencarian teori (Bab 2), pemilihan metode (Bab 3), "
        "pelaksanaan analisis dan pembahasan berbasis teori (Bab 4), hingga "
        "sintesis kesimpulan yang menjawab rumusan masalah secara sirkular (Bab 5)."
    )
    intro_run.font.size = Pt(11)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    # Ekstrak bagian tabel simpul dari output agen
    simpul_raw = ""
    tabel_m = re.search(
        r'\[TABEL SIMPUL\](.*?)\[/TABEL SIMPUL\]',
        audit_text, re.DOTALL | re.IGNORECASE
    )
    if tabel_m:
        simpul_raw = tabel_m.group(1)

    # Render tabel 5 SIMPUL — selalu tampil
    _add_simpul_table(doc, simpul_raw)

    # Ekstrak & render Matriks Keterlacakan RM-Tujuan-Hasil-Kesimpulan
    matrix_m = re.search(
        r'\[MATRIKS KETERLACAKAN\](.*?)\[/MATRIKS KETERLACAKAN\]',
        audit_text, re.DOTALL | re.IGNORECASE
    )
    if matrix_m:
        _add_traceability_matrix(doc, matrix_m.group(1))

    # Narasi audit (teks sebelum marker [SOLUSI REVISI])
    full_consistency = audit_text  # fallback
    narasi_m = re.search(
        r'(.*?)\[SOLUSI REVISI — CONSISTENCY ENGINE\]',
        audit_text, re.DOTALL | re.IGNORECASE
    )
    if narasi_m:
        narasi_text = _strip_machine_blocks(narasi_m.group(1).strip())
        if narasi_text:
            _add_heading(doc, "Narasi Audit", 2)
            _render_rich_text(doc, narasi_text)
    doc.add_paragraph()

    # Blok revisi per simpul
    consistency_section = sections.get("consistency", "")
    if consistency_section:
        _add_heading(doc, "Temuan & Solusi Revisi per Simpul", 2)
        # Blok data mesin [TABEL SIMPUL] & [MATRIKS KETERLACAKAN] sudah dirender
        # jadi tabel di atas; kalau agen menaruhnya di dalam bagian SOLUSI REVISI,
        # buang teks mentahnya supaya baris "SIMPUL 1 | ... | TERHUBUNG" / "RM1 |
        # ... | TERJAWAB" tidak bocor muncul dua kali sebagai paragraf jelek.
        _add_simpul_revision_blocks(doc, _strip_machine_blocks(consistency_section))
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 2: Statistical Auditor ─────────────────────────────────────
    _add_heading(doc, "2. Audit Statistik & Metodologi (Bab 3–4)", 1)
    _add_perspective_label(doc, "Metodologi, Statistik & Ketepatan Angka")
    if sections.get("stats"):
        _render_rich_text(doc, sections["stats"])
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 3: Discussion Critique ─────────────────────────────────────
    _add_heading(doc, "3. Kritik Bab Pembahasan", 1)
    _add_perspective_label(doc, "Argumentasi, Logika & Kebijakan")
    if sections.get("discussion"):
        _render_rich_text(doc, sections["discussion"])
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 4: Ghost Citation ──────────────────────────────────────────
    _add_heading(doc, "4. Deteksi Referensi Hantu & Referensi Yatim", 1)
    _add_perspective_label(doc, "Kepatuhan Akademik & Integritas Sitasi")
    if sections.get("ghost"):
        _render_rich_text(doc, sections["ghost"])
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 5: Paraphrasing Engine ────────────────────────────────────
    _add_heading(doc, "5. Parafrase Kalimat Berindikasi Plagiarisme", 1)
    if sections.get("paraphrase"):
        _render_rich_text(doc, sections["paraphrase"])
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 6: Journal Compatibility ──────────────────────────────────
    _add_heading(doc, "6. Penilaian Kelayakan Scopus Q1", 1)
    journal_text = sections.get("journal", "")
    # Buang isi Draf Revisi Penuh dari sini kalau ikut tertangkap (dirender
    # terpisah di Bagian 7) supaya tidak duplikat.
    journal_text = re.split(r'\[DRAF REVISI PENUH — BAHASA INDONESIA\]', journal_text, flags=re.IGNORECASE)[0].strip()
    if journal_text:
        _render_rich_text(doc, journal_text)
    else:
        doc.add_paragraph("(Analisis sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 7: Draf Revisi Penuh (Bahasa Indonesia) ───────────────────
    _add_heading(doc, "7. Draf Revisi Penuh — Bahasa Indonesia (Siap Tempel)", 1)
    notice_id = doc.add_paragraph()
    notice_id_run = notice_id.add_run(
        "CATATAN: Draf ini menggabungkan semua perbaikan yang ditemukan tim audit. "
        "WAJIB diverifikasi ulang oleh mahasiswa dan dosen pembimbing sebelum "
        "digunakan, terutama angka-angka hasil penelitian."
    )
    notice_id_run.font.italic = True
    notice_id_run.font.color.rgb = RGBColor(0x85, 0x4F, 0x0B)
    notice_id_run.font.size = Pt(10)
    if sections.get("draf_revisi_id"):
        _render_rich_text(doc, sections["draf_revisi_id"])
    else:
        doc.add_paragraph("(Draf revisi penuh tidak tersedia untuk analisis ini.)")
    doc.add_page_break()

    # ── Bagian 8: Auto-Template Converter ────────────────────────────────
    _add_heading(doc, "8. Draf Artikel Scopus Q1 — Siap Kirim", 1)
    notice = doc.add_paragraph()
    notice_run = notice.add_run(
        "CATATAN: Teks di bawah ini adalah draf yang dihasilkan oleh RMD-TRA "
        "berdasarkan isi skripsi Anda. Lakukan verifikasi, penyesuaian, dan "
        "pemformatan sesuai author guidelines jurnal target sebelum pengiriman."
    )
    notice_run.font.italic = True
    notice_run.font.color.rgb = RGBColor(0x85, 0x4F, 0x0B)
    notice_run.font.size = Pt(10)

    if sections.get("template"):
        _render_rich_text(doc, sections["template"])
    else:
        doc.add_paragraph("(Template sedang diproses atau tidak tersedia.)")
    doc.add_page_break()

    # ── Bagian 9: Simulasi Sidang ─────────────────────────────────────────
    _add_heading(doc, "9. Simulasi Pertanyaan Sidang & Jawaban Aman", 1)
    notice_sidang = doc.add_paragraph()
    notice_sidang_run = notice_sidang.add_run(
        "Daftar ini disusun berdasarkan seluruh temuan audit di atas, diurutkan "
        "dari yang paling berisiko ditanyakan penguji. Latih jawabannya, jangan "
        "dihafal kata-per-kata — pahami intinya."
    )
    notice_sidang_run.font.italic = True
    notice_sidang_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    notice_sidang_run.font.size = Pt(10)
    if sections.get("sidang"):
        _render_rich_text(doc, sections["sidang"])
    else:
        doc.add_paragraph("(Simulasi sidang tidak tersedia untuk analisis ini.)")

    # ── Footer ─────────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"RMD-TRA v2.0 — Laporan dihasilkan otomatis pada "
        f"{datetime.now().strftime('%d %B %Y, %H:%M')} WIB"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    # ── Simpan ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    if output_path:
        with open(output_path, "wb") as f:
            f.write(content)

    return content
