"""
RMD-TRA — tests/test_word_exporter.py
Test parsing & rendering tabel di laporan Word Pilar 5:
1. Tabel 5 SIMPUL — regresi untuk bug lama: status (TERHUBUNG/LEMAH/TERPUTUS)
   tidak pernah terisi karena sub_agent_consistency belum diinstruksikan
   menghasilkan format [TABEL SIMPUL] yang benar.
2. Matriks Keterlacakan RM-Tujuan-Hasil-Kesimpulan (fitur baru).

Jalankan: uv run pytest tests/test_word_exporter.py -v
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import docx  # noqa: E402

from app.pillar5.word_exporter import generate_audit_report  # noqa: E402


def _table_rows_as_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def test_simpul_table_status_is_parsed_correctly():
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
SIMPUL 2 | Bab 2 -> Bab 3 | Teori -> Metode | [STATUS: LEMAH]
SIMPUL 3 | Bab 3 -> Bab 4 | Metode -> Hasil | [STATUS: TERHUBUNG]
SIMPUL 4 | Bab 4 -> Bab 5 | Hasil -> Kesimpulan | [STATUS: TERHUBUNG]
SIMPUL 5 | Bab 1 <-> Bab 5 | Sirkularitas | [STATUS: TERPUTUS]
[/TABEL SIMPUL]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]
"""
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa")
    import io
    doc = docx.Document(io.BytesIO(docx_bytes))

    simpul_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "Simpul")
    rows = _table_rows_as_text(simpul_table)
    statuses = {row[0]: row[3] for row in rows[1:]}
    assert statuses["SIMPUL 1"] == "TERHUBUNG"
    assert statuses["SIMPUL 2"] == "LEMAH"
    assert statuses["SIMPUL 5"] == "TERPUTUS"
    # Regresi: sebelum diperbaiki, semua baris jatuh ke fallback "—"
    assert "—" not in statuses.values()


def test_traceability_matrix_is_rendered():
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
SIMPUL 2 | Bab 2 -> Bab 3 | Teori -> Metode | [STATUS: TERHUBUNG]
SIMPUL 3 | Bab 3 -> Bab 4 | Metode -> Hasil | [STATUS: TERHUBUNG]
SIMPUL 4 | Bab 4 -> Bab 5 | Hasil -> Kesimpulan | [STATUS: TERHUBUNG]
SIMPUL 5 | Bab 1 <-> Bab 5 | Sirkularitas | [STATUS: TERHUBUNG]
[/TABEL SIMPUL]

[MATRIKS KETERLACAKAN]
RM1 | Bagaimana kebutuhan energi? | Mendeskripsikan kebutuhan energi | 4.4 Analisis Energi | Terpenuhi | [STATUS: LENGKAP]
RM2 | Bagaimana biaya operasional? | Menganalisis biaya | TIDAK DITEMUKAN | TIDAK DITEMUKAN | [STATUS: TIDAK LENGKAP]
[/MATRIKS KETERLACAKAN]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]
"""
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa")
    import io
    doc = docx.Document(io.BytesIO(docx_bytes))

    matrix_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "RM")
    rows = _table_rows_as_text(matrix_table)
    assert rows[1][0] == "RM1"
    assert rows[1][5] == "LENGKAP"
    assert rows[2][3] == "TIDAK DITEMUKAN"
    assert rows[2][5] == "TIDAK LENGKAP"


def test_missing_traceability_matrix_does_not_crash():
    """Kalau agen tidak menghasilkan blok matriks, laporan tetap ter-generate tanpa error."""
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
[/TABEL SIMPUL]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]
"""
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa")
    assert len(docx_bytes) > 0


def test_finance_verification_section_rendered_when_provided():
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
[/TABEL SIMPUL]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]
"""
    fv = "Status konsistensi angka: KONSISTEN — tidak ditemukan angka yang bertentangan"
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa", finance_verification=fv)
    import io
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Verifikasi Angka Otomatis" in full_text
    assert "KONSISTEN" in full_text


def test_sidang_prep_section_is_parsed_and_rendered():
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
[/TABEL SIMPUL]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]

[SIMULASI SIDANG — PERTANYAAN & JAWABAN]
KATEGORI: Pertanyaan Umum
TINGKAT RISIKO: Sedang
PERTANYAAN PENGUJI: "Kenapa memilih lokasi penelitian ini?"
JAWABAN AMAN: "Karena lokasi ini representatif untuk kasus yang diteliti."
[SIDANG_PREP_DONE]
"""
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa")
    import io
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Simulasi Pertanyaan Sidang" in full_text
    assert "Kenapa memilih lokasi penelitian ini?" in full_text


def test_draf_revisi_indonesia_split_from_journal_section():
    """
    Bagian 'Draf Revisi Penuh Bahasa Indonesia' harus dirender terpisah dari
    Bagian 6 (Penilaian Scopus Q1), bukan ikut tercampur/duplikat di sana.
    """
    audit_text = """
[TABEL SIMPUL]
SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | [STATUS: TERHUBUNG]
[/TABEL SIMPUL]

[SOLUSI REVISI — CONSISTENCY ENGINE]
Contoh.
[CONSISTENCY_DONE]

[DRAF IMRAD — SIAP TEMPEL DAN DIKEMBANGKAN]
TITLE: Contoh Judul Artikel Ilmiah

[DRAF REVISI PENUH — BAHASA INDONESIA]
ABSTRAK
Ini konten abstrak revisi bahasa Indonesia yang unik.
[JOURNAL_DONE]
"""
    docx_bytes = generate_audit_report(audit_text, "Uji", "Mahasiswa")
    import io
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Draf Revisi Penuh" in full_text
    assert "Ini konten abstrak revisi bahasa Indonesia yang unik." in full_text
    # Pastikan tidak duplikat: cuma muncul SEKALI, bukan 2x (sekali di Bagian 6, sekali di 7)
    assert full_text.count("Ini konten abstrak revisi bahasa Indonesia yang unik.") == 1
