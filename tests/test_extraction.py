"""
RMD-TRA — tests/test_extraction.py
Regression test untuk dua bug ekstraksi dokumen Pilar 5 yang pernah ditemukan:

1. _extract_chapters salah menangkap ringkasan "Sistematika Penulisan" (yang
   menyebut ulang BAB I..V secara singkat & berdekatan) sebagai bab sungguhan,
   sehingga Bab II-V yang sebenarnya (jauh lebih panjang) tidak terbaca sama
   sekali. Ditemukan dari audit nyata skripsi Abim & Variz (090626).

2. _extract_from_docx melewatkan seluruh isi tabel Word (cuma baca
   doc.paragraphs), sehingga tabel data finansial (NPV/IRR/BCR) hilang total
   dan dianggap "KOSONG" oleh sub-agen audit.

Jalankan: uv run pytest tests/test_extraction.py -v
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from app.fast_api_app import _extract_chapters, _extract_from_docx  # noqa: E402
from app.pillar5.finance_tools import (  # noqa: E402
    extract_labeled_figures,
    recompute_bca_from_tables,
)


# ---------------------------------------------------------------------------
# Bug #1 — Ekstraksi bab tertipu ringkasan "Sistematika Penulisan"
# ---------------------------------------------------------------------------
def _build_thesis_with_toc_trap() -> str:
    """
    Rekonstruksi pola nyata: Bab I asli, lalu ringkasan sistematika penulisan
    yang menyebut BAB I..V berdekatan (jarak antar penyebutan < 800 karakter),
    baru kemudian bab-bab asli yang panjang.
    """
    real_bab2 = "BAB II TINJAUAN LITERATUR\n\n" + ("Konten tinjauan pustaka. " * 200)
    real_bab3 = "BAB III METODE PENELITIAN\n\n" + ("Konten metodologi penelitian. " * 200)
    real_bab4 = "BAB IV HASIL DAN PEMBAHASAN\n\n" + ("Konten hasil dan pembahasan. " * 200)
    real_bab5 = "BAB V KESIMPULAN DAN SARAN\n\n" + ("Konten kesimpulan. " * 200)

    sistematika = (
        "BAB I Pendahuluan\nBerisi latar belakang, rumusan masalah, tujuan.\n\n"
        "BAB II Tinjauan Pustaka\nBerisi teori dan penelitian terdahulu.\n\n"
        "BAB III Metode Penelitian\nMenjelaskan metode dan teknik analisis.\n\n"
        "BAB IV Hasil dan Pembahasan\nMenyajikan hasil dan pembahasan.\n\n"
        "BAB V Kesimpulan dan Saran\nBerisi kesimpulan dan saran.\n\n"
    )

    text = (
        "BAB I\nPENDAHULUAN\n\n" + ("Konten latar belakang penelitian. " * 100)
        + "\n\n" + sistematika
        + real_bab2 + "\n\n" + real_bab3 + "\n\n" + real_bab4 + "\n\n" + real_bab5
        + "\n\nDAFTAR PUSTAKA\n\nPenulis, A. (2023). Judul referensi. Jurnal X.\n"
    )
    return text


def test_extract_chapters_ignores_sistematika_penulisan_summary():
    text = _build_thesis_with_toc_trap()
    chapters = _extract_chapters(text)

    # Bab II-V harus menangkap isi ASLI (ratusan karakter "Konten ..." berulang),
    # bukan cuma cuplikan ringkasan sistematika (~50-80 karakter).
    assert "TINJAUAN LITERATUR" in chapters["bab2"]
    assert chapters["bab2"].count("Konten tinjauan pustaka.") > 50
    assert len(chapters["bab2"]) > 1000

    assert "METODE PENELITIAN" in chapters["bab3"]
    assert chapters["bab3"].count("Konten metodologi penelitian.") > 50
    assert len(chapters["bab3"]) > 1000

    assert "HASIL DAN PEMBAHASAN" in chapters["bab4"]
    assert chapters["bab4"].count("Konten hasil dan pembahasan.") > 50
    assert len(chapters["bab4"]) > 1000

    assert "KESIMPULAN DAN SARAN" in chapters["bab5"]
    assert chapters["bab5"].count("Konten kesimpulan.") > 50
    # Bab 5 tidak boleh "bocor" berisi konten bab lain
    assert "TINJAUAN LITERATUR" not in chapters["bab5"]

    assert "Judul referensi" in chapters["dapus"]


def test_extract_chapters_missing_chapter_returns_empty_not_wrong_content():
    text = "BAB I\nPENDAHULUAN\n\n" + ("Konten. " * 100)
    chapters = _extract_chapters(text)
    assert chapters["bab2"] == ""
    assert chapters["bab3"] == ""


def test_extract_chapters_keeps_short_but_genuine_chapter():
    """
    Edge case: bab yang cuma disebut SEKALI di seluruh dokumen (tidak ada
    kandidat lain untuk dibandingkan) harus tetap dipakai meski isinya pendek
    (<MIN_GAP karakter) — mis. Daftar Pustaka singkat. Filter jarak hanya boleh
    aktif kalau memang ada >1 kandidat untuk bab yang sama.
    """
    text = (
        "BAB I\nPENDAHULUAN\n\n" + ("Konten pendahuluan. " * 100)
        + "\n\nBAB II TINJAUAN LITERATUR\n\n" + ("Konten tinjauan. " * 100)
        + "\n\nDAFTAR PUSTAKA\nPenulis, A. (2023). Judul singkat.\n"
    )
    chapters = _extract_chapters(text)
    assert "Judul singkat" in chapters["dapus"]
    assert chapters["dapus"] != ""


def test_extract_chapters_strips_sistematika_penulisan_from_content():
    """
    Bukan cuma dihindari saat mendeteksi titik mulai bab — isi ringkasan
    "Sistematika Penulisan" (yang menyebut ulang BAB I..V berdekatan) harus
    benar-benar dibuang dari hasil akhir Bab 1, tidak ikut terkirim ke agen.
    """
    text = _build_thesis_with_toc_trap()
    chapters = _extract_chapters(text)

    assert "Berisi teori dan penelitian terdahulu" not in chapters["bab1"]
    assert "Menjelaskan metode dan teknik analisis" not in chapters["bab1"]
    assert "Menyajikan hasil dan pembahasan" not in chapters["bab1"]
    assert "Berisi kesimpulan dan saran" not in chapters["bab1"]
    # Konten asli Bab 1 tetap utuh, tidak ikut terpotong
    assert chapters["bab1"].count("Konten latar belakang penelitian.") > 50


def test_extract_chapters_keeps_genuine_cross_reference_mid_sentence():
    """
    Referensi silang di tengah kalimat (mis. "...sebagaimana dijelaskan pada
    Bab IV, bus listrik memiliki...") BUKAN bagian dari sistematika penulisan
    — harus tetap dipertahankan, bukan ikut terbuang sebagai decoy.
    """
    real_bab4 = (
        "BAB IV HASIL DAN PEMBAHASAN\n\n"
        + ("Konten pembahasan awal. " * 150)
        + "\nSebagaimana dijelaskan pada Bab IV, bus listrik memiliki biaya operasional lebih rendah.\n"
        + ("Konten pembahasan lanjutan. " * 150)
    )
    text = (
        "BAB I\nPENDAHULUAN\n\n" + ("Konten. " * 50)
        + "\n\nBAB II TINJAUAN LITERATUR\n\n" + ("Konten. " * 150)
        + "\n\nBAB III METODE\n\n" + ("Konten. " * 150)
        + "\n\n" + real_bab4
        + "\n\nBAB V KESIMPULAN\n\n" + ("Konten. " * 100)
        + "\n\nDAFTAR PUSTAKA\nPenulis, A. (2023). Judul.\n"
    )
    chapters = _extract_chapters(text)
    assert "sebagaimana dijelaskan pada Bab IV, bus listrik memiliki".lower() in chapters["bab4"].lower()
    assert chapters["bab4"].count("Konten pembahasan awal.") > 100
    assert chapters["bab4"].count("Konten pembahasan lanjutan.") > 100


# ---------------------------------------------------------------------------
# Bug #2 — Ekstraksi DOCX melewatkan tabel
# ---------------------------------------------------------------------------
def _build_docx_with_table() -> bytes:
    import docx

    doc = docx.Document()
    doc.add_paragraph("Paragraf sebelum tabel.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Nilai"
    table.rows[1].cells[0].text = "NPV"
    table.rows[1].cells[1].text = "Rp 10.329.812.880"
    doc.add_paragraph("Paragraf setelah tabel.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_extract_from_docx_includes_table_content():
    content = _build_docx_with_table()
    result = await _extract_from_docx(content)
    text = result["text"]

    assert "Paragraf sebelum tabel." in text
    assert "Paragraf setelah tabel." in text
    assert "[TABEL]" in text and "[/TABEL]" in text
    assert "NPV" in text
    assert "Rp 10.329.812.880" in text

    # Urutan dokumen harus terjaga: sebelum -> tabel -> sesudah
    idx_before = text.index("Paragraf sebelum tabel.")
    idx_table = text.index("[TABEL]")
    idx_after = text.index("Paragraf setelah tabel.")
    assert idx_before < idx_table < idx_after


# ---------------------------------------------------------------------------
# Bug #3 — Verifikasi angka salah atribusi subjek di tabel perbandingan
# ---------------------------------------------------------------------------
def test_finance_no_false_inconsistency_from_comparison_table():
    """
    Regresi: di tabel perbandingan dengan header berkelompok
    ("Diesel|Diesel|Listrik|Listrik"), nilai kolom diesel JANGAN sampai salah
    dilabeli sebagai bus listrik. Sekarang tabel diparse COLUMN-AWARE (subjek
    dari posisi kolom), jadi nilai tabel yang konsisten dengan narasi tidak
    memicu false positive. Ditemukan dari audit nyata skripsi Abim.
    """
    text = (
        "Hasil analisis menunjukkan nilai NPV bus listrik sebesar Rp 4.596.212.966 "
        "dan IRR sebesar 30,8%. Adapun NPV bus diesel sebesar Rp 10.329.812.880 "
        "dengan IRR sebesar 109,9%.\n\n"
        "[TABEL]\n"
        "BCA | Bus Diesel | Bus Diesel | Bus Listrik | Bus Listrik\n"
        "NPV | 10.329.812.880 | layak | 4.596.212.966 | layak\n"
        "IRR | 109,9% | layak | 30,8% | layak\n"
        "[/TABEL]\n"
    )
    fig = extract_labeled_figures(text)
    assert fig["inconsistencies_found"] == 0, (
        f"False positive: {[i['warning'] for i in fig['inconsistencies']]}"
    )


def test_finance_detects_real_narrative_inconsistency():
    """Deteksi TETAP bekerja untuk inkonsistensi sungguhan di narasi
    (mis. abstrak vs Bab IV menyebut NPV bus listrik berbeda)."""
    text = (
        "Abstrak: nilai NPV bus listrik sebesar Rp 5.290.190.887.\n"
        "Bab IV: hasil menunjukkan NPV bus listrik sebesar Rp 4.596.212.966.\n"
    )
    fig = extract_labeled_figures(text)
    assert fig["inconsistencies_found"] >= 1


def test_finance_column_aware_catches_table_vs_narrative_mismatch():
    """
    Kemampuan baru column-aware: kalau nilai di TABEL berbeda dengan narasi untuk
    subjek yang sama, itu HARUS tertangkap (dulu tabel dikecualikan → terlewat)."""
    text = (
        "Abstrak menyebut NPV bus listrik sebesar Rp 5.290.190.887.\n\n"
        "[TABEL]\n"
        "BCA | Bus Diesel | Bus Listrik\n"
        "NPV | 10.329.812.880 | 4.596.212.966\n"
        "[/TABEL]\n"
    )
    fig = extract_labeled_figures(text)
    # NPV bus listrik: narasi 5,29M vs tabel 4,6M -> inkonsisten
    npv_listrik = [
        i for i in fig["inconsistencies"]
        if i["indicator"] == "NPV" and i["subject"] == "bus listrik"
    ]
    assert npv_listrik, "Mismatch tabel-vs-narasi untuk NPV bus listrik tidak tertangkap"


def test_recompute_single_column_not_forced():
    """Tabel satu kolom 'Net Cash' (ambigu) TIDAK boleh dipaksa jadi NPV."""
    text = (
        "[TABEL]\n"
        "Tahun | Net Cash | Discount Rate | PV\n"
        "1 | 1.477.505.000 | 0.9091 | 1.343.186.364\n"
        "2 | 1.481.505.000 | 0.8264 | 1.224.384.298\n"
        "[/TABEL]\n"
    )
    res = recompute_bca_from_tables(text)
    assert res["recomputed"], "tabel arus kas tidak terdeteksi sama sekali"
    assert all(not r.get("reliable") for r in res["recomputed"]), (
        "tabel satu-kolom seharusnya ditandai tidak-andal, bukan dihitung sebagai NPV"
    )


def test_recompute_net_columns_reliable():
    """Tabel dengan kolom Manfaat DAN Biaya -> NPV/BCR dihitung andal."""
    text = (
        "[TABEL]\n"
        "Tahun | Manfaat | Biaya\n"
        "1 | 2.000.000.000 | 1.500.000.000\n"
        "2 | 2.000.000.000 | 1.500.000.000\n"
        "3 | 2.000.000.000 | 1.500.000.000\n"
        "[/TABEL]\n"
    )
    res = recompute_bca_from_tables(text)
    reliable = [r for r in res["recomputed"] if r.get("reliable")]
    assert reliable, "tabel manfaat+biaya seharusnya bisa dihitung andal"
    assert reliable[0]["bcr_recomputed"] is not None and reliable[0]["bcr_recomputed"] > 1


# ---------------------------------------------------------------------------
# Bug #5 — Parser section terlalu kaku: header agen sedikit beda -> section hilang
# ---------------------------------------------------------------------------
def test_parse_sections_tolerates_marker_variations():
    """
    Regresi: agen kadang menulis header dengan tanda hubung berbeda ("-" bukan
    "—") atau tanpa kurung siku. Dulu itu membuat SELURUH section hilang dari
    laporan Word (kasus nyata: Bagian 8 template ~15rb kata lenyap). Parser
    harus toleran, dan kalau header benar-benar tak dikenali tapi marker
    penutup [X_DONE] ada, konten tetap terselamatkan lewat fallback.
    """
    from app.pillar5.word_exporter import _parse_sections

    audit = (
        "[SOLUSI REVISI - CONSISTENCY ENGINE]\n"       # dash biasa, bukan em-dash
        "isi konsistensi di sini\n"
        "[CONSISTENCY_DONE]\n"
        "HEADER TEMPLATE YANG TIDAK DIKENALI SAMA SEKALI\n"  # header rusak total
        "=== VERSI A: TEMPLATE 1 KOLOM (IEEE) ===\n"
        "isi artikel lengkap dua versi\n"
        "[TEMPLATE_DONE]\n"
    )
    sections = _parse_sections(audit)
    assert "isi konsistensi" in sections["consistency"], "variasi dash harus tetap dikenali"
    assert "isi artikel lengkap" in sections["template"], (
        "header rusak: konten wajib terselamatkan lewat fallback marker penutup"
    )


def test_parse_sections_recovers_when_end_marker_missing():
    """
    Regresi (kasus nyata run 4 Jul): agen template menulis header dengan benar
    tapi LUPA marker penutup [TEMPLATE_DONE] setelah output panjang — Bagian 8
    (~15rb kata) lenyap dari laporan. Fallback A: header ada + penutup hilang
    -> ambil sampai marker section berikutnya / EOF.
    """
    from app.pillar5.word_exporter import _parse_sections

    audit = (
        "[PARAPHRASE_DONE]"
        "[AUTO-TEMPLATE CONVERTER — SIAP KIRIM KE JURNAL]\n"
        "=== VERSI A: TEMPLATE 1 KOLOM (IEEE) ===\n"
        "isi artikel versi IEEE yang sangat panjang\n"
        "=== VERSI B: TEMPLATE 2 KOLOM (APA 7th) ===\n"
        "isi artikel versi APA\n"
        # TIDAK ada [TEMPLATE_DONE] — langsung section berikutnya
        "[SIMULASI SIDANG — PERTANYAAN & JAWABAN]\n"
        "KATEGORI: Metodologi\nPERTANYAAN PENGUJI: contoh\n"
        "[SIDANG_PREP_DONE]\n"
    )
    sections = _parse_sections(audit)
    assert "isi artikel versi IEEE" in sections["template"]
    assert "isi artikel versi APA" in sections["template"]
    # konten sidang tidak boleh bocor ke section template
    assert "PERTANYAAN PENGUJI" not in sections["template"]
    assert "PERTANYAAN PENGUJI" in sections["sidang"]


# ---------------------------------------------------------------------------
# Bug #6 — Tabel merged-cell dari Word asli tampil berantakan di laporan
# ---------------------------------------------------------------------------
def test_render_table_cleans_merged_cell_artifacts():
    """
    Regresi (kasus nyata Gambar 4.1/4.2 pemetaan bus & Tabel 4.4 komparasi):
    sel gabungan Word terbaca berulang ("Depo Bus | Depo Bus", "D=19,4 km" x5),
    baris header terduplikasi, dan ada kolom kosong total — tampil tak
    terbaca. Renderer wajib: buang baris duplikat berurutan, buang kolom
    kosong, dan menggabungkan sel identik berurutan jadi satu sel.
    """
    from docx import Document
    from app.pillar5.word_exporter import _render_table_from_rows

    rows = [
        ["Depo Bus", "Depo Bus", "D=19,4 km", "D=19,4 km", "D=19,4 km", "Halte", ""],
        ["Depo Bus", "Depo Bus", "D=19,4 km", "D=19,4 km", "D=19,4 km", "Halte", ""],  # duplikat
        ["1", "", "2", "3", "4", "5", ""],
    ]
    d = Document()
    _render_table_from_rows(d, rows)
    t = d.tables[0]

    # baris duplikat hilang: 3 baris input -> 2 baris tabel
    assert len(t.rows) == 2, f"baris duplikat harus dibuang, dapat {len(t.rows)}"
    # kolom kosong total (terakhir) hilang: 7 -> 6 kolom
    assert len(t.columns) == 6, f"kolom kosong harus dibuang, dapat {len(t.columns)}"
    # sel identik berurutan digabung: "Depo Bus" x2 -> satu sel (tc sama)
    hdr = t.rows[0]
    assert hdr.cells[0]._tc is hdr.cells[1]._tc, "sel 'Depo Bus' ganda harus tergabung"
    # "D=19,4 km" x3 -> satu sel
    assert hdr.cells[2]._tc is hdr.cells[4]._tc, "sel 'D=19,4 km' x3 harus tergabung"
    # baris data tidak ikut tergabung salah (nilai beda tetap terpisah)
    data = t.rows[1]
    assert data.cells[2]._tc is not data.cells[3]._tc


def test_render_table_compacts_sparse_specification_table():
    """
    Regresi (kasus nyata Tabel 4.4 Komparasi Biaya Operasional): agen jurnal
    mereproduksi tabel spesifikasi label:nilai dari skripsi asli sebagai
    tabel markdown 15 kolom dengan header terduplikasi 15x dan tiap baris
    data hanya berisi 1-3 sel terisi tersebar di kolom yang berbeda-beda
    ("| | | Kapasitas Battery | 357 | Kwh | | | ... |"). Membuang kolom yang
    kosong di SEMUA baris tidak menolong di sini karena tiap kolom terisi
    di baris yang berbeda -> tabel tampil lebar & kosong ("berantakan").
    Renderer wajib merapatkan sel berisi tiap baris ke kiri dan memangkas
    kolom ke jumlah minimum yang benar-benar dipakai.
    """
    from docx import Document
    from app.pillar5.word_exporter import _render_table_from_rows

    title = "ELECTRIC BUS (EB): VKTR BYD D9"
    ncol = 15
    rows = [
        [title] * ncol,
        ["Spesifikasi :"] + [""] * (ncol - 1),
        ["", "", "Kapasitas Battery", "357", "Kwh"] + [""] * (ncol - 5),
        ["", "", "Effisiensi Energi", "", "85%-90%"] + [""] * (ncol - 5),
        ["Nilai Investasi (CAPEX)", "Nilai Investasi (CAPEX)", "Rp", "4,500,000,000"] + [""] * (ncol - 4),
    ]
    d = Document()
    _render_table_from_rows(d, rows)
    t = d.tables[0]

    # dirapatkan jauh dari 15 kolom asli -> maksimal 3 kolom (label,nilai,satuan)
    assert len(t.columns) <= 4, f"tabel sparse harus dipangkas, dapat {len(t.columns)} kolom"
    body_text = [[c.text for c in r.cells] for r in t.rows]
    flat = [cell for row in body_text for cell in row]
    # data yang tadinya tersebar sekarang harus tetap ada & terbaca (tidak hilang)
    assert "Kapasitas Battery" in flat
    assert "357" in flat
    assert "Kwh" in flat
    assert "4,500,000,000" in flat
    # tidak ada baris yang hampir seluruhnya kosong (kepadatan tinggi lagi)
    for row in body_text[1:]:
        filled = sum(1 for c in row if c.strip())
        assert filled >= 1, f"baris data tidak boleh kosong total: {row}"


# ---------------------------------------------------------------------------
# Bug #4 — Kurung kurawal LaTeX di instruction dikira variabel state ADK
# ---------------------------------------------------------------------------
def test_agent_instructions_bypass_state_injection():
    """
    Regresi: ADK memperlakukan {nama} di instruction STRING sebagai variabel
    session-state dan melempar KeyError kalau tidak ada. Instruksi agen kita
    mengandung kurung kurawal literal (rumus LaTeX \\frac{a}{b}) — pernah
    membuat 2 agen paralel crash ("unhandled errors in a TaskGroup").
    Semua instruction Pilar 5 WAJIB berupa callable (InstructionProvider)
    supaya state-injection dilewati, dan teksnya harus tetap utuh.
    """
    from app.pillar5 import agents as a

    all_agents = [
        a.sub_agent_consistency, a.sub_agent_stats_auditor,
        a.sub_agent_discussion, a.sub_agent_ghost_citation,
        a.sub_agent_journal, a.sub_agent_paraphrase,
        a.sub_agent_template, a.sub_agent_sidang_prep,
    ]
    for ag in all_agents:
        assert callable(ag.instruction), (
            f"{ag.name}: instruction masih string — kurung kurawal LaTeX akan "
            "dikira variabel state ADK dan bikin agen crash"
        )
        text = ag.instruction(None)
        assert isinstance(text, str) and len(text) > 100, (
            f"{ag.name}: teks instruction hilang/terpotong setelah dibungkus"
        )


# ---------------------------------------------------------------------------
# Perubahan #1 — Guardrail ekstraksi bab (kegagalan diam-diam -> bersuara)
# ---------------------------------------------------------------------------
def test_guardrail_warns_when_one_chapter_swallows_document():
    """Kalau satu bab menelan porsi tak wajar, guardrail harus berteriak."""
    from app.fast_api_app import _diagnose_chapter_extraction

    chapters = {
        "bab1": "x" * 100, "bab2": "x" * 100, "bab3": "x" * 100,
        "bab4": "y" * 9000, "bab5": "x" * 100, "dapus": "x" * 50,
    }
    total = sum(len(v) for v in chapters.values())
    warns = _diagnose_chapter_extraction(chapters, total)
    assert any("menempati" in w for w in warns), (
        "Bab yang menelan >72% dokumen harus memicu peringatan"
    )


def test_guardrail_silent_on_normal_thesis():
    """Skripsi normal (bab terbagi wajar) TIDAK boleh memunculkan peringatan palsu."""
    from app.fast_api_app import _diagnose_chapter_extraction

    chapters = {
        "bab1": "a" * 8000, "bab2": "b" * 40000, "bab3": "c" * 18000,
        "bab4": "d" * 42000, "bab5": "e" * 5000, "dapus": "f" * 3000,
    }
    total = sum(len(v) for v in chapters.values())
    warns = _diagnose_chapter_extraction(chapters, total)
    assert warns == [], f"Skripsi normal tidak boleh diperingatkan, dapat: {warns}"


# ---------------------------------------------------------------------------
# Perubahan #2 — Verifikator survei (Slovin) untuk skripsi non-finansial
# ---------------------------------------------------------------------------
def test_slovin_detects_margin_contradiction():
    """
    Pola nyata (skripsi stasiun Duri): margin ditulis 5% tapi sampel 100
    sebenarnya pakai 10%. Slovin harus menangkap kontradiksi ini secara pasti.
    """
    from app.pillar5.survey_tools import verify_slovin

    text = (
        "Jumlah sampel ditentukan menggunakan Rumus Slovin dengan batas toleransi "
        "kesalahan (error margin) sebesar 5%. Populasi penelitian adalah seluruh "
        "pengguna jasa dengan rata rata 28.554 penumpang/hari. Berdasarkan hasil "
        "perhitungan dengan rumus Slovin pada tingkat toleransi kesalahan sebesar "
        "10%, diperoleh kebutuhan jumlah responden minimal sebanyak 100 responden."
    )
    r = verify_slovin(text)
    assert r["applicable"] is True
    assert r["population"] == 28554
    assert 0.05 in r["margins_mentioned"] and 0.1 in r["margins_mentioned"]
    assert r["sample_stated"] == 100
    assert r["issues"], "Kontradiksi margin 5% vs 10% harus ditandai"


def test_slovin_silent_when_not_used():
    """Skripsi finansial (tanpa Slovin) tidak boleh memicu temuan survei palsu."""
    from app.pillar5.survey_tools import verify_slovin

    text = "Analisis kelayakan memakai NPV, IRR, dan BCR atas arus kas 20 tahun."
    r = verify_slovin(text)
    assert r["applicable"] is False and r["issues"] == []


def test_survey_block_empty_for_financial_thesis():
    """
    Blok verifikasi survei harus KOSONG untuk skripsi finansial murni, supaya
    laporan tidak diganggu bagian tak relevan (netralisasi celah gagal-diam).
    """
    from app.fast_api_app import _build_survey_verification_block

    text = "Studi kelayakan elektrifikasi bus. NPV Rp 5 miliar, IRR 12%, BCR 1,3."
    assert _build_survey_verification_block(text) == ""


def test_finance_block_suppressed_for_nonfinancial_thesis():
    """
    Skripsi survei (tanpa materi finansial) TIDAK boleh memunculkan blok
    verifikasi finansial — bahkan untuk bilang "tidak ada tabel arus kas".
    Kata "BCA"/NPV milik skripsi finansial; menampilkannya di skripsi survei
    adalah kebocoran konteks antar-skripsi.
    """
    from app.fast_api_app import _build_finance_verification_block

    survey_text = (
        "Penelitian kepuasan pengguna dengan kuesioner skala Likert di stasiun. "
        "Analisis memakai CSI dan IPA berdasarkan SPM PM 63/2019. Rumus Slovin "
        "dipakai untuk menentukan 100 responden dari populasi penumpang."
    )
    assert _build_finance_verification_block(survey_text) == ""


def test_finance_block_present_for_financial_thesis():
    """Skripsi finansial WAJIB tetap memunculkan blok verifikasi (jangan ikut mati)."""
    from app.fast_api_app import _build_finance_verification_block

    fin_text = (
        "Analisis kelayakan investasi memakai NPV, IRR, dan BCR. "
        "[TABEL]\nTahun | Manfaat | Biaya\n0 | 0 | 1000\n1 | 600 | 100\n"
        "2 | 600 | 100\n3 | 600 | 100\n[/TABEL] Discount rate 10%."
    )
    block = _build_finance_verification_block(fin_text)
    assert block and ("NPV" in block or "BCA" in block or "hitung ulang" in block.lower())


def test_survey_block_not_triggered_by_incidental_spm_citation():
    """
    Skripsi finansial yang cuma MENYITASI "SPM" sekali (mis. "Standar Pelayanan
    Minimal (SPM) PT Transjakarta (Setiawan, 2020)") TIDAK boleh memunculkan blok
    survei — itu kebocoran materi survei ke laporan finansial. Butuh penanda
    metodologi survei nyata (kuesioner/Slovin/CSI/IPA/responden), bukan sitasi.
    """
    from app.fast_api_app import _build_survey_verification_block

    fin_text = (
        "Analisis kelayakan investasi bus listrik memakai NPV, IRR, dan BCR atas "
        "arus kas 20 tahun. Penelitian mengacu pada Standar Pelayanan Minimal (SPM) "
        "PT Transjakarta (Setiawan, 2020). Diharapkan hasil memiliki validitas dan "
        "reliabilitas yang baik."
    )
    assert _build_survey_verification_block(fin_text) == ""


def test_survey_block_present_for_survey_thesis():
    """Untuk skripsi survei, blok harus terisi dan memuat hasil Slovin."""
    from app.fast_api_app import _build_survey_verification_block

    text = (
        "Penelitian kepuasan pengguna dengan kuesioner skala Likert. Uji validitas "
        "dan reliabilitas dilakukan. Rumus Slovin dengan toleransi kesalahan 10% "
        "atas populasi rata rata 28.554 penumpang/hari menghasilkan 100 responden. "
        "Metode CSI dan IPA (importance performance analysis) dipakai untuk SPM."
    )
    block = _build_survey_verification_block(text)
    assert block and "SLOVIN" in block and "100" in block


def test_strip_machine_blocks_removes_raw_simpul_and_matrix_lines():
    """
    Regresi: blok [TABEL SIMPUL] & [MATRIKS KETERLACAKAN] dirender jadi TABEL,
    tapi teks mentahnya ("SIMPUL 1 | ... | TERHUBUNG", "RM1 | ... | TERJAWAB")
    pernah bocor muncul lagi sebagai paragraf jelek. _strip_machine_blocks harus
    membuangnya total dari narasi/section, TAPI mempertahankan prosa biasa.
    """
    from app.pillar5.word_exporter import _strip_machine_blocks

    raw = (
        "Narasi audit yang mengalir tentang benang merah skripsi ini.\n"
        "[TABEL SIMPUL]\n"
        "SIMPUL 1 | Bab 1 -> Bab 2 | Masalah -> Teori | TERHUBUNG\n"
        "SIMPUL 2 | Bab 2 -> Bab 3 | Teori -> Metode | LEMAH\n"
        "[/TABEL SIMPUL]\n"
        "Paragraf penilaian keseluruhan yang harus tetap ada.\n"
        "[MATRIKS KETERLACAKAN]\n"
        "RM1 | pertanyaan | tujuan | jawaban | kesimpulan | TERJAWAB\n"
        "[/MATRIKS KETERLACAKAN]\n"
        "Kalimat penutup narasi.\n"
        "RM2 | baris mesin tercecer di luar blok | x | y | z | PERLU DICEK"
    )
    out = _strip_machine_blocks(raw)
    assert "SIMPUL 1 |" not in out and "SIMPUL 2 |" not in out
    assert "RM1 |" not in out and "RM2 |" not in out
    assert "[TABEL SIMPUL]" not in out and "[MATRIKS KETERLACAKAN]" not in out
    # prosa biasa harus lestari
    assert "Narasi audit yang mengalir" in out
    assert "Paragraf penilaian keseluruhan" in out
    assert "Kalimat penutup narasi." in out
